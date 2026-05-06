import asyncio
import logging
from dataclasses import dataclass
from app.config import get_settings

logger = logging.getLogger(__name__)


@dataclass
class ParseResult:
    raw_text: str
    markdown: str
    metadata: dict


async def parse_document(file_path: str, file_type: str) -> ParseResult:
    """根据文件类型分发到对应解析器"""
    parsers = {
        "pdf": _parse_pdf,
        "docx": _parse_docx,
        "md": _parse_markdown,
    }
    parser = parsers.get(file_type)
    if not parser:
        raise ValueError(f"不支持的文件类型: {file_type}")

    return await parser(file_path)


async def _parse_pdf(file_path: str) -> ParseResult:
    """使用 pymupdf4llm 解析 PDF 为 Markdown"""
    import pymupdf4llm

    # pymupdf4llm 是同步的，放到线程池中运行
    loop = asyncio.get_event_loop()
    md_text = await loop.run_in_executor(None, pymupdf4llm.to_markdown, file_path)

    import pymupdf
    doc = pymupdf.open(file_path)
    raw_text = ""
    for page in doc:
        raw_text += page.get_text()
    metadata = {
        "pages": len(doc),
        "title": doc.metadata.get("title", ""),
    }
    doc.close()

    return ParseResult(raw_text=raw_text, markdown=md_text, metadata=metadata)


async def _parse_docx(file_path: str) -> ParseResult:
    """使用 python-docx 解析 DOCX"""
    from docx import Document as DocxDocument

    loop = asyncio.get_event_loop()

    def _extract():
        doc = DocxDocument(file_path)
        paragraphs = []
        raw_parts = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            raw_parts.append(text)

            # 根据样式转换为 Markdown
            style_name = para.style.name.lower() if para.style else ""
            if "heading 1" in style_name:
                paragraphs.append(f"# {text}")
            elif "heading 2" in style_name:
                paragraphs.append(f"## {text}")
            elif "heading 3" in style_name:
                paragraphs.append(f"### {text}")
            elif "heading 4" in style_name:
                paragraphs.append(f"#### {text}")
            else:
                paragraphs.append(text)

        # 处理表格
        for table in doc.tables:
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append("| " + " | ".join(cells) + " |")
            if rows:
                # 在表头后添加分隔行
                header = rows[0]
                separator = "| " + " | ".join(["---"] * len(table.rows[0].cells)) + " |"
                paragraphs.append(header)
                paragraphs.append(separator)
                paragraphs.extend(rows[1:])

        return "\n".join(raw_parts), "\n\n".join(paragraphs)

    raw_text, md_text = await loop.run_in_executor(None, _extract)

    return ParseResult(
        raw_text=raw_text,
        markdown=md_text,
        metadata={"format": "docx"},
    )


async def _parse_markdown(file_path: str) -> ParseResult:
    """直接读取 Markdown 文件"""
    with open(file_path, "r", encoding="utf-8") as f:
        content = f.read()

    return ParseResult(
        raw_text=content,
        markdown=content,
        metadata={"format": "markdown"},
    )
